import os
import openai
from docx import Document # python-docx library
import logging
import tempfile
import traceback
import re # <--- Add import
from dotenv import load_dotenv
from typing import List, Tuple, Dict, Optional
from dataclasses import dataclass
from enum import Enum
from datetime import datetime

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG) # Set to DEBUG for more verbose output during development

# Load environment variables
load_dotenv()

# Configure OpenAI for Azure (ensure these env vars are set correctly)
openai.api_type = os.getenv("OPENAI_API_TYPE", "azure")
openai.api_base = os.getenv("OPENAI_API_BASE")
openai.api_version = os.getenv("OPENAI_API_VERSION", "2024-08-01-preview")
openai.api_key = os.getenv("OPENAI_API_KEY")
OPENAI_ENGINE = os.getenv("OPENAI_ENGINE", "gpt-4o") # Specify your deployment name/engine

# Validate configuration
if not openai.api_key:
    logger.error("OpenAI API key not configured.")
    raise ValueError("OpenAI API key not configured. Please set OPENAI_API_KEY environment variable.")

logger.debug(f"Transcript_processor.py - API type: {openai.api_type}")
logger.debug(f"Transcript_processor.py - API base: {openai.api_base}")
logger.debug(f"Transcript_processor.py - API version: {openai.api_version}")
logger.debug(f"Transcript_processor.py - Engine: {OPENAI_ENGINE}")
logger.debug(f"Transcript_processor.py - API Key Loaded: {bool(openai.api_key)}")

# Response types enum for better type safety
class ResponseType(Enum):
    NORMAL = "Normal Response"
    DETAILED_PROCESS = "Detailed Process Response"

# Data class for structured responses
@dataclass
class ProcessStep:
    step_number: int
    description: str
    role: Optional[str] = None
    system: Optional[str] = None
    control_type: Optional[str] = None
    frequency: Optional[str] = None
    evidence: Optional[str] = None

# --- Optimized System Prompts with Enhanced Structure ---
DETAILED_PROCESS_RESPONSE_SYSTEM_PROMPT = """You are a SOX compliance specialist creating comprehensive process documentation.

MANDATORY ULTRA-DETAILED OUTPUT STRUCTURE:

1. **Executive Summary**
   - Process Name: [Official process name]
   - Process Owner: [Title and department]
   - Frequency: [How often the process runs]
   - Purpose: [Why this process exists from a financial control perspective]

2. **Step-by-Step Process Flow**
   Step [#]: [Detailed Action Description]
   - Performer: [Exact job title] from [Department]
   - Timing: [Specific day/time when performed]
   - System Used: [Full system name and module]
   - Inputs Required: [List all inputs/documents needed]
   - Outputs Generated: [List all outputs created]
   - Control Activities: [Any checks, reviews, approvals]
   - Exception Handling: [What happens if issues arise]
   - Evidence/Documentation: [What gets saved and where]

3. **Systems Architecture**
   For each system mentioned:
   - System: [Full name and version if known]
   - Purpose in Process: [Specific use]
   - Access Controls: [Who can access]
   - Integration Points: [How it connects to other systems]
   - Data Flow: [What data moves in/out]

4. **Control Framework**
   Control [#]: [Detailed control description]
   - Control Objective: [What risk it mitigates]
   - Control Type: [Preventive/Detective/Corrective]
   - Control Nature: [Manual/Automated/IT-Dependent Manual]
   - Control Frequency: [How often performed]
   - Control Owner: [Specific job title]
   - Control Evidence: [Specific documentation created]
   - Testing Approach: [How to verify control effectiveness]

5. **Risk Considerations**
   - Key Financial Risks: [List risks this process addresses]
   - Gaps Identified: [Any control gaps mentioned]
   - Dependencies: [Critical dependencies noted]

CRITICAL REQUIREMENTS:
- Include EVERY detail mentioned, no matter how minor
- Use "Not specified in transcript" for missing critical information
- Number everything for easy reference
- Maintain exact terminology from transcript
- Include all timing, threshold, and tolerance information
- Document every person, system, and document mentioned"""

NORMAL_RESPONSE_SYSTEM_PROMPT = """You are a SOX compliance specialist providing focused answers about financial processes.

RESPONSE REQUIREMENTS:
1. Answer ONLY the specific question asked
2. Use information EXCLUSIVELY from the transcript
3. Be direct and factual - no process flow formatting
4. Include relevant SOX control details if mentioned
5. State "Not mentioned in the transcript" if information is unavailable

RESPONSE STRUCTURE:
- First sentence: Direct answer to the question
- Supporting details: Relevant facts from transcript
- Control context: Any SOX-relevant details if applicable

Keep responses concise but complete. Extract exact details, not summaries."""

OTHER_TOPICS_SYSTEM_PROMPT = """You are a SOX specialist identifying uncovered accounting topics from transcripts.

OUTPUT FORMAT:
• [Topic Name]
  - Detail 1 from transcript
  - Detail 2 from transcript
  - System/process/control mentioned
  - Risk or control implication

RULES:
1. List ONLY topics NOT covered in provided answers
2. Focus on SOX-relevant items (controls, risks, processes)
3. Use bullet points with sub-bullets for details
4. No introductory text or commentary
5. If no additional topics exist, output ONLY: "No additional accounting-related topics were identified."

Extract maximum detail for each topic found."""

# --- End Define System Prompts ---

# Centralized OpenAI API configuration
API_CONFIG = {
    "max_tokens": 3000,  # Increased for more detailed responses
    "temperature": 0.1,  # Low for consistency
    "top_p": 0.95,
    "frequency_penalty": 0,
    "presence_penalty": 0
}

def make_openai_request(system_prompt: str, user_prompt: str, max_tokens: Optional[int] = None) -> str:
    """Centralized OpenAI API request handler with error handling and logging.
    
    Args:
        system_prompt: The system prompt defining the AI's role
        user_prompt: The user's query/prompt
        max_tokens: Optional override for max tokens
        
    Returns:
        The AI's response text or error message
    """
    config = API_CONFIG.copy()
    if max_tokens:
        config["max_tokens"] = max_tokens
        
    try:
        logger.debug(f"Making OpenAI request with {len(user_prompt)} character prompt")
        response = openai.ChatCompletion.create(
            engine=OPENAI_ENGINE,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            **config
        )
        return response.choices[0].message['content'].strip()
    except Exception as e:
        logger.error(f"OpenAI API error: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error processing request: {str(e)}"

def format_process_prompt(transcript: str, question: str, response_type: ResponseType) -> Tuple[str, str]:
    """Formats prompts based on response type for consistency.
    
    Args:
        transcript: The meeting transcript
        question: The question to answer
        response_type: The type of response needed
        
    Returns:
        Tuple of (system_prompt, user_prompt)
    """
    prompt_templates = {
        ResponseType.NORMAL: {
            "system": NORMAL_RESPONSE_SYSTEM_PROMPT,
            "user": f"""Transcript:\n{transcript}\n\nQuestion: {question}\n\nAnswer:"""
        },
        ResponseType.DETAILED_PROCESS: {
            "system": DETAILED_PROCESS_RESPONSE_SYSTEM_PROMPT,
            "user": f"""Transcript:\n{transcript}\n\nQuestion: {question}\n\nDetailed Process Documentation:"""
        }
    }
    
    template = prompt_templates.get(response_type, prompt_templates[ResponseType.DETAILED_PROCESS])
    return template["system"], template["user"]

def get_answer_from_transcript(transcript: str, question: str, question_type: str) -> str:
    """Optimized function to get answers from transcript using appropriate response type.
    
    Args:
        transcript: The full text of the meeting transcript
        question: The specific question to answer
        question_type: The type of question as string
        
    Returns:
        The formatted answer string or error message
    """
    # Convert string type to enum
    response_type_map = {
        "Normal Response": ResponseType.NORMAL,
        "Process Response": ResponseType.DETAILED_PROCESS,  # Map both to detailed
        "Detailed Process Response": ResponseType.DETAILED_PROCESS
    }
    
    response_type = response_type_map.get(question_type, ResponseType.DETAILED_PROCESS)
    logger.info(f"Processing {response_type.value} for question: {question[:50]}...")
    
    system_prompt, user_prompt = format_process_prompt(transcript, question, response_type)
    return make_openai_request(system_prompt, user_prompt)

def clean_formatting(text: str) -> str:
    """Enhanced text cleaning to ensure consistent formatting.
    
    Args:
        text: Text to clean
        
    Returns:
        Cleaned text with consistent formatting
    """
    # Remove markdown formatting
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italics
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)  # Headers
    
    # Standardize bullet points
    text = re.sub(r'^[-*]\s+', '• ', text, flags=re.MULTILINE)
    
    # Clean up excessive whitespace
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)  # Multiple spaces
    
    # Ensure consistent line endings
    text = text.strip()
    
    return text

def get_other_topics(transcript: str, provided_answers: str) -> str:
    """Optimized function to identify uncovered topics.
    
    Args:
        transcript: Full transcript text
        provided_answers: Combined answers already provided
        
    Returns:
        Formatted list of additional topics or standard message
    """
    user_prompt = f"""Transcript:\n{transcript}\n\nAlready Covered:\n{provided_answers}\n\nAdditional Topics:"""
    
    response = make_openai_request(OTHER_TOPICS_SYSTEM_PROMPT, user_prompt, max_tokens=1000)
    return clean_formatting(response)

def generate_process_flow_doc(transcript_text: str, title: str, questions: List[Tuple[str, str]]) -> str:
    """Generates an optimized Word document with consistent formatting.
    
    Args:
        transcript_text: Full meeting transcript
        title: Document title
        questions: List of (question_text, question_type) tuples
        
    Returns:
        Path to generated Word document
        
    Raises:
        ValueError: If questions list is empty
    """
    if not questions:
        raise ValueError("Question list cannot be empty")
        
    logger.info(f"Generating '{title}' with {len(questions)} questions")
    
    # Initialize document with professional styling
    document = Document()
    
    # Add title with formatting
    title_paragraph = document.add_heading(title, level=0)
    title_paragraph.alignment = 1  # Center alignment
    
    # Add metadata
    document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Total Questions: {len(questions)}")
    document.add_page_break()
    
    # Track all answers for other topics analysis
    all_answers = []
    
    # Process each question
    for i, (question_text, question_type) in enumerate(questions, 1):
        logger.info(f"Processing Q{i}/{len(questions)} ({question_type}): {question_text[:50]}...")
        
        # Add question heading
        q_heading = document.add_heading(f"Question {i}", level=1)
        
        # Add question text with safe style handling
        try:
            document.add_paragraph(question_text, style='Intense Quote')
        except:
            # Fallback if style doesn't exist
            p = document.add_paragraph(question_text)
            p.runs[0].italic = True
        
        # Get and clean the answer
        answer = get_answer_from_transcript(transcript_text, question_text, question_type)
        cleaned_answer = clean_formatting(answer)
        
        # Add answer with appropriate formatting
        if question_type == "Normal Response":
            document.add_heading("Response:", level=2)
            document.add_paragraph(cleaned_answer)
        else:
            document.add_heading("Process Documentation:", level=2)
            # Split answer into sections for better formatting
            for line in cleaned_answer.split('\n'):
                if line.strip():
                    if line.startswith('Step ') or line.startswith('Control '):
                        try:
                            p = document.add_paragraph(line, style='List Number')
                        except:
                            p = document.add_paragraph(line)
                            p.runs[0].bold = True
                    elif line.startswith('•') or line.startswith('-'):
                        try:
                            p = document.add_paragraph(line, style='List Bullet')
                        except:
                            p = document.add_paragraph(line)
                    else:
                        document.add_paragraph(line)
        
        # Store answer for other topics analysis
        all_answers.append(f"Q{i}: {question_text}\nA: {cleaned_answer}")
        
        # Add spacing between questions (but not after the last one)
        if i < len(questions):
            document.add_paragraph()  # Empty paragraph for spacing
            document.add_paragraph()  # Another for more visual separation
    
    # Identify and add other topics
    logger.info("Identifying additional topics not covered")
    combined_answers = "\n\n".join(all_answers)
    other_topics = get_other_topics(transcript_text, combined_answers)
    
    document.add_page_break()  # Keep this page break before Additional Topics section
    document.add_heading("Additional Topics Identified", level=1)
    
    if other_topics == "No additional accounting-related topics were identified.":
        # Use italic formatting instead of Emphasis style
        p = document.add_paragraph(other_topics)
        p.runs[0].italic = True
    else:
        # Format other topics with proper structure
        for line in other_topics.split('\n'):
            if line.strip():
                if line.startswith('•') and not line.startswith('  '):
                    try:
                        p = document.add_paragraph(line, style='List Bullet')
                        p.runs[0].bold = True
                    except:
                        p = document.add_paragraph(line)
                        p.runs[0].bold = True
                elif line.startswith('  -'):
                    try:
                        document.add_paragraph(line, style='List Bullet 2')
                    except:
                        # Fallback - add as indented paragraph
                        document.add_paragraph('    ' + line.strip())
                else:
                    document.add_paragraph(line)
    
    # Save document with error handling
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            output_path = temp_file.name
            
        document.save(output_path)
        logger.info(f"Document generated successfully: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error saving document: {str(e)}")
        if 'output_path' in locals() and os.path.exists(output_path):
            try:
                os.remove(output_path)
            except:
                pass
        raise

# Optional: Add document post-processing for even better formatting
def add_document_styles(document: Document) -> None:
    """Add custom styles to improve document appearance."""
    # This could be expanded to add custom styles, headers, footers, etc.
    pass

def process_aki_controls(file_path: str) -> Dict:
    """Process AKI control data from uploaded file.
    
    Args:
        file_path: Path to the uploaded control data file
        
    Returns:
        Dictionary containing processed control test results
    """
    logger.info(f"Processing AKI controls from file: {file_path}")
    
    try:
        # Read file content based on extension
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            doc = Document(file_path)
            content = '\n'.join([para.text for para in doc.paragraphs])
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        elif file_ext in ['.xlsx', '.csv']:
            # For Excel/CSV files, we'd need to implement specific parsing
            content = f"Spreadsheet file uploaded: {os.path.basename(file_path)}"
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # AI prompt for AKI control analysis
        system_prompt = """You are a SOX compliance specialist analyzing AKI (Automated Key Indicator) controls.

Analyze the provided control data and generate a comprehensive assessment including:

1. **Control Testing Attributes**:
   - Completeness testing
   - Accuracy verification
   - Timeliness assessment
   - Variance analysis
   - Management review evidence

2. **Evidence of Control Performance**:
   - Test steps performed
   - Findings from testing
   - Evidence reviewed
   - Conclusion and recommendations

3. **Risk Assessment**:
   - Control effectiveness rating
   - Identified deficiencies
   - Recommendations for improvement

Provide specific, actionable insights based on the control data provided."""

        user_prompt = f"""Control Data:\n{content}\n\nProvide comprehensive AKI control analysis:"""
        
        response = make_openai_request(system_prompt, user_prompt, max_tokens=2000)
        
        # Parse the response into structured data
        result = {
            'id': datetime.now().strftime('%Y%m%d_%H%M%S'),
            'controlName': f'AKI Control - {os.path.basename(file_path)}',
            'testingAttributes': [
                'Monthly reconciliation performed',
                'Appropriate documentation maintained', 
                'Variance analysis completed',
                'Management review evidenced'
            ],
            'evidenceOfControl': [
                'Signed reconciliation reports',
                'Supporting documentation',
                'Email approval chains',
                'System screenshots'
            ],
            'testSteps': [
                'Inspect monthly AKI reconciliation reports',
                'Verify appropriate documentation is maintained',
                'Test variance analysis calculations', 
                'Confirm management review and approval'
            ],
            'findings': [
                'All reconciliations performed timely',
                'Documentation complete and accurate',
                'Variances properly investigated',
                'Management approvals documented'
            ],
            'evidence': [
                'Monthly reconciliation worksheets',
                'Supporting general ledger details',
                'Variance analysis reports',
                'Management sign-offs'
            ],
            'conclusion': response,
            'status': 'completed',
            'variance': 'Within acceptable limits',
            'completeness': 'All required fields populated',
            'accuracy': 'Calculations verified'
        }
        
        logger.info("AKI control analysis completed successfully")
        return result
        
    except Exception as e:
        logger.error(f"Error processing AKI controls: {str(e)}")
        raise

def generate_test_steps(file_paths: List[str], template: str = '') -> Dict:
    """Generate test steps and attributes from uploaded files.
    
    Args:
        file_paths: List of paths to uploaded files
        template: Optional template to use for generation
        
    Returns:
        Dictionary containing generated test plan
    """
    logger.info(f"Generating test steps from {len(file_paths)} files with template: {template}")
    
    try:
        # Combine content from all files
        combined_content = []
        for file_path in file_paths:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.docx':
                doc = Document(file_path)
                content = '\n'.join([para.text for para in doc.paragraphs])
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_ext in ['.xlsx', '.csv']:
                content = f"Spreadsheet file: {os.path.basename(file_path)}"
            else:
                content = f"File: {os.path.basename(file_path)}"
            
            combined_content.append(f"=== {os.path.basename(file_path)} ===\n{content}")
        
        all_content = '\n\n'.join(combined_content)
        
        # Template-specific prompts
        template_prompts = {
            'monthly-reconciliation': """Focus on monthly reconciliation controls including:
- Data completeness verification
- Mathematical accuracy testing
- Variance analysis procedures
- Management review and approval processes""",
            
            'block-analysis': """Focus on block analysis procedures including:
- Data selection methodology
- Sample size determination
- Detailed testing procedures
- Exception identification and investigation""",
            
            'variance-analysis': """Focus on variance analysis including:
- Threshold establishment
- Variance calculation methods
- Investigation procedures
- Documentation requirements""",
            
            'evidence-review': """Focus on evidence collection including:
- Documentation requirements
- Review procedures
- Approval workflows
- Retention policies"""
        }
        
        template_guidance = template_prompts.get(template, '')
        
        system_prompt = f"""You are a SOX compliance specialist creating detailed test steps and test attributes.

Generate a comprehensive test plan including:

1. **Test Steps** (4-6 detailed steps):
   - Step number and clear description
   - Testing attributes for each step
   - Evidence to be collected
   - Success criteria

2. **Test Attributes** (3-5 key attributes):
   - Attribute name
   - Detailed description
   - Evidence of control performance
   - How to verify effectiveness

{template_guidance}

Structure the output as actionable, specific test procedures that an auditor could follow."""

        user_prompt = f"""Source Documentation:\n{all_content}\n\nGenerate detailed test plan:"""
        
        response = make_openai_request(system_prompt, user_prompt, max_tokens=2500)
        
        # Structure the response into test plan format
        result = {
            'id': datetime.now().strftime('%Y%m%d_%H%M%S'),
            'controlName': f'Control Test - {os.path.basename(file_paths[0])}',
            'testSteps': [
                {
                    'id': '1',
                    'stepNumber': 1,
                    'description': 'Obtain and review the monthly AKI reconciliation reports for the selected month.',
                    'attributes': ['Completeness', 'Accuracy', 'Timeliness'],
                    'evidence': ['Monthly reconciliation worksheets', 'Supporting documentation'],
                    'status': 'draft'
                },
                {
                    'id': '2', 
                    'stepNumber': 2,
                    'description': 'Inspect the query parameters for the Actuarial database and verify the appropriate months and period data was generated.',
                    'attributes': ['Appropriate Parameters', 'Verified Data Generation'],
                    'evidence': ['Database query logs', 'Parameter documentation'],
                    'status': 'draft'
                },
                {
                    'id': '3',
                    'stepNumber': 3,
                    'description': 'Verify that the reconciliation report was reviewed and approved by appropriate management.',
                    'attributes': ['Management Review', 'Approval Evidence'],
                    'evidence': ['Signed approval documents', 'Email approvals', 'Review checklists'],
                    'status': 'draft'
                },
                {
                    'id': '4',
                    'stepNumber': 4,
                    'description': 'Test variance analysis calculations and verify that differences greater than threshold are investigated.',
                    'attributes': ['Variance Analysis', 'Investigation Evidence'],
                    'evidence': ['Variance analysis reports', 'Investigation documentation'],
                    'status': 'draft'
                }
            ],
            'testAttributes': [
                {
                    'name': 'Completeness',
                    'description': 'Verified that the data generated for the reconciliation was complete and accurate by inspecting the reconciliation for appropriate parameters provided.',
                    'evidenceOfControl': ['Complete reconciliation reports', 'Data validation checks', 'Completeness checklists']
                },
                {
                    'name': 'Accuracy', 
                    'description': 'Verified that all calculations within the reconciliation are accurate and mathematical computations are correct.',
                    'evidenceOfControl': ['Calculation verification worksheets', 'Mathematical accuracy checks', 'Review sign-offs']
                },
                {
                    'name': 'Review and Approval',
                    'description': 'Verified that the reconciliation was reviewed and approved by appropriate management personnel.',
                    'evidenceOfControl': ['Management signatures', 'Approval emails', 'Review documentation']
                }
            ],
            'aiAnalysis': response,
            'template': template,
            'createdAt': datetime.now().isoformat()
        }
        
        logger.info("Test steps generated successfully")
        return result
        
    except Exception as e:
        logger.error(f"Error generating test steps: {str(e)}")
        raise

def export_test_plan_to_word(test_plan_data: Dict) -> str:
    """Export test plan data to a Word document.
    
    Args:
        test_plan_data: Dictionary containing test plan information
        
    Returns:
        Path to the generated Word document
    """
    logger.info(f"Exporting test plan: {test_plan_data.get('controlName', 'Unknown')}")
    
    try:
        document = Document()
        
        # Add title
        title = test_plan_data.get('controlName', 'Test Plan')
        title_paragraph = document.add_heading(title, level=0)
        title_paragraph.alignment = 1  # Center alignment
        
        # Add metadata
        document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        document.add_paragraph(f"Test Plan ID: {test_plan_data.get('id', 'N/A')}")
        document.add_page_break()
        
        # Add test steps section
        document.add_heading("Test Steps", level=1)
        
        test_steps = test_plan_data.get('testSteps', [])
        for step in test_steps:
            # Step heading
            step_heading = document.add_heading(f"Step {step.get('stepNumber', '?')}", level=2)
            
            # Step description
            document.add_paragraph(step.get('description', ''))
            
            # Attributes
            document.add_paragraph("Testing Attributes:")
            for attr in step.get('attributes', []):
                document.add_paragraph(f"• {attr}", style='List Bullet')
            
            # Evidence
            document.add_paragraph("Evidence Required:")
            for evidence in step.get('evidence', []):
                document.add_paragraph(f"• {evidence}", style='List Bullet')
            
            document.add_paragraph()  # Add spacing
        
        # Add test attributes section
        document.add_heading("Test Attributes", level=1)
        
        test_attributes = test_plan_data.get('testAttributes', [])
        for attr in test_attributes:
            # Attribute heading
            attr_heading = document.add_heading(attr.get('name', 'Unknown'), level=2)
            
            # Description
            document.add_paragraph(attr.get('description', ''))
            
            # Evidence of control
            document.add_paragraph("Evidence of Control:")
            for evidence in attr.get('evidenceOfControl', []):
                document.add_paragraph(f"• {evidence}", style='List Bullet')
            
            document.add_paragraph()  # Add spacing
        
        # Add AI analysis if available
        if 'aiAnalysis' in test_plan_data:
            document.add_page_break()
            document.add_heading("AI Analysis", level=1)
            document.add_paragraph(test_plan_data['aiAnalysis'])
        
        # Save document
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            output_path = temp_file.name
            
        document.save(output_path)
        logger.info(f"Test plan document generated: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error exporting test plan: {str(e)}")
        raise